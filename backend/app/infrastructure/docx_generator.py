"""DOCX generator for corrected contracts.

Generates formatted Word documents with corrected contract text
and a summary of all corrections made.
"""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from app.infrastructure.gemini_models import CorrectedContractResult


def generate_corrected_contract_docx(
    result: CorrectedContractResult,
    contract_title: str = "Contrato Corrigido",
) -> BytesIO:
    """Generate a DOCX file with the corrected contract and corrections summary.

    Args:
        result: CorrectedContractResult from Gemini analysis.
        contract_title: Title for the document header.

    Returns:
        BytesIO buffer containing the DOCX file.
    """
    doc = Document()
    
    # Set up styles
    _setup_styles(doc)
    
    # Header with LegalBoard branding
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("LegalBoard")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 188, 212)  # Cyan
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(contract_title)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Spacing
    
    # Corrected Contract Section
    section_title = doc.add_paragraph()
    run = section_title.add_run("CONTRATO CORRIGIDO")
    run.bold = True
    run.font.size = Pt(12)
    
    # Add corrected contract text
    for paragraph_text in result.corrected_text.split("\n\n"):
        if paragraph_text.strip():
            para = doc.add_paragraph(paragraph_text.strip())
            para.paragraph_format.space_after = Pt(6)
    
    # Page break before corrections summary
    doc.add_page_break()
    
    # Corrections Summary Section
    summary_title = doc.add_paragraph()
    run = summary_title.add_run("RESUMO DAS CORREÇÕES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 188, 212)
    
    doc.add_paragraph()
    
    if not result.corrections:
        doc.add_paragraph("Nenhuma correção foi necessária.")
    else:
        for i, correction in enumerate(result.corrections, 1):
            # Correction header
            corr_header = doc.add_paragraph()
            run = corr_header.add_run(f"{i}. {correction.clause_name}")
            run.bold = True
            run.font.size = Pt(11)
            
            # Reason
            reason_para = doc.add_paragraph()
            run = reason_para.add_run("Motivo: ")
            run.bold = True
            reason_para.add_run(correction.reason)
            
            # Original text (truncated if too long)
            original_text = correction.original_text
            if len(original_text) > 300:
                original_text = original_text[:300] + "..."
            
            orig_para = doc.add_paragraph()
            run = orig_para.add_run("Texto original: ")
            run.bold = True
            run.font.color.rgb = RGBColor(180, 0, 0)  # Red
            orig_para.add_run(original_text)
            
            # Corrected text (truncated if too long)
            corrected_text = correction.corrected_text
            if len(corrected_text) > 300:
                corrected_text = corrected_text[:300] + "..."
            
            corr_para = doc.add_paragraph()
            run = corr_para.add_run("Texto corrigido: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            corr_para.add_run(corrected_text)
            
            doc.add_paragraph()  # Spacing between corrections
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Gerado por LegalBoard • Plataforma de Governança Contratual")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def _setup_styles(doc: Document) -> None:
    """Set up document styles."""
    # Modify Normal style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
