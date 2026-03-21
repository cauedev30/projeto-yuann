"""Tests for DOCX generator."""

from io import BytesIO

import pytest
from docx import Document

from app.infrastructure.docx_generator import generate_corrected_contract_docx
from app.infrastructure.gemini_models import CorrectedContractResult, CorrectionItem


class TestGenerateCorrectedContractDocx:
    def test_generates_valid_docx(self):
        result = CorrectedContractResult(
            corrected_text="Este é o contrato corrigido.\n\nCláusula 1: Texto.",
            corrections=[
                CorrectionItem(
                    clause_name="INFRAESTRUTURA",
                    original_text="Texto antigo",
                    corrected_text="Texto novo corrigido",
                    reason="Adequação ao playbook",
                )
            ],
        )

        buffer = generate_corrected_contract_docx(result)

        assert isinstance(buffer, BytesIO)
        assert buffer.tell() == 0  # Should be at start

        # Verify it's a valid DOCX
        doc = Document(buffer)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        
        assert "LegalBoard" in full_text
        assert "CONTRATO CORRIGIDO" in full_text
        assert "RESUMO DAS CORREÇÕES" in full_text
        assert "INFRAESTRUTURA" in full_text

    def test_handles_empty_corrections(self):
        result = CorrectedContractResult(
            corrected_text="Contrato sem correções necessárias.",
            corrections=[],
        )

        buffer = generate_corrected_contract_docx(result)
        doc = Document(buffer)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Nenhuma correção foi necessária" in full_text

    def test_uses_custom_title(self):
        result = CorrectedContractResult(
            corrected_text="Texto",
            corrections=[],
        )

        buffer = generate_corrected_contract_docx(
            result,
            contract_title="Contrato de Locação Comercial",
        )
        doc = Document(buffer)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Contrato de Locação Comercial" in full_text

    def test_truncates_long_text_in_corrections(self):
        long_text = "A" * 500  # Longer than 300 char limit
        result = CorrectedContractResult(
            corrected_text="Texto",
            corrections=[
                CorrectionItem(
                    clause_name="TEST",
                    original_text=long_text,
                    corrected_text=long_text,
                    reason="Test",
                )
            ],
        )

        buffer = generate_corrected_contract_docx(result)
        doc = Document(buffer)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Should be truncated with "..."
        assert "..." in full_text
        assert long_text not in full_text  # Full text should not appear
